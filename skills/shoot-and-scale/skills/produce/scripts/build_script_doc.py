# -*- coding: utf-8 -*-
"""
Shoot & Scale — Script Doc builder.
Usage:  python build_script_doc.py <data.json> <output.docx>

data.json schema:
{
  "shoot": "Shoot 2",
  "client": "Devin Bender",
  "ig": "https://www.instagram.com/devbender/",
  "videos": [
    {
      "topic": "....",
      "format": "Says vs Means",
      "format_link": "https://...",          # optional; renders as a 'HERE' hyperlink
      "text_hook": "....",                    # optional
      "editor_notes": "....",                 # optional; usually blank
      "script": [                              # optional; list of bullet items
        {"who": "client", "t": "..."},        # who: client (black) | interviewer (red)
        {"runs": [                             # OR mixed colors on one bullet
            {"who": "interviewer", "t": "Question?"},
            {"who": "client", "t": " Answer."}
        ]}
      ]
    }
  ]
}
"""
import sys, os, json
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SKILL_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
ASSETS = os.path.join(SKILL_DIR, "assets")
LOGO = os.path.join(ASSETS, "logo.png")
COMMENT = os.path.join(ASSETS, "comment_help.png")

RED=RGBColor(0xFF,0,0); BLACK=RGBColor(0,0,0); LABELGRAY=RGBColor(0x33,0x33,0x33)
LABEL_FILL="EDEDED"; NOTE_FILL="F4F7FB"; BORDER="C7C7C7"; LINKBLUE="0563C1"

def shade(cell,hexc):
    tcPr=cell._tc.get_or_add_tcPr(); shd=OxmlElement('w:shd')
    shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto'); shd.set(qn('w:fill'),hexc); tcPr.append(shd)
def borders(cell,color=BORDER,sz='6',edges=('top','left','bottom','right')):
    tcPr=cell._tc.get_or_add_tcPr(); tb=OxmlElement('w:tcBorders')
    for e in ('top','left','bottom','right'):
        el=OxmlElement(f'w:{e}')
        if e in edges:
            el.set(qn('w:val'),'single'); el.set(qn('w:sz'),sz); el.set(qn('w:space'),'0'); el.set(qn('w:color'),color)
        else: el.set(qn('w:val'),'nil')
        tb.append(el)
    tcPr.append(tb)
def set_w(cell,inches):
    cell.width=Inches(inches); tcPr=cell._tc.get_or_add_tcPr(); w=OxmlElement('w:tcW')
    w.set(qn('w:w'),str(int(inches*1440))); w.set(qn('w:type'),'dxa'); tcPr.append(w)
def fix_table(table,widths):
    tbl=table._tbl; tblPr=tbl.tblPr
    layout=OxmlElement('w:tblLayout'); layout.set(qn('w:type'),'fixed'); tblPr.append(layout)
    tblW=OxmlElement('w:tblW'); tblW.set(qn('w:w'),str(int(sum(widths)*1440))); tblW.set(qn('w:type'),'dxa'); tblPr.append(tblW)
    g=tbl.find(qn('w:tblGrid'))
    if g is not None: tbl.remove(g)
    g=OxmlElement('w:tblGrid')
    for wi in widths:
        gc=OxmlElement('w:gridCol'); gc.set(qn('w:w'),str(int(wi*1440))); g.append(gc)
    tblPr.addnext(g); table.allow_autofit=False
def add_hyperlink(paragraph, url, text):
    part=paragraph.part
    r_id=part.relate_to(url,"http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",is_external=True)
    h=OxmlElement('w:hyperlink'); h.set(qn('r:id'),r_id)
    r=OxmlElement('w:r'); rPr=OxmlElement('w:rPr')
    c=OxmlElement('w:color'); c.set(qn('w:val'),LINKBLUE); rPr.append(c)
    u=OxmlElement('w:u'); u.set(qn('w:val'),'single'); rPr.append(u)
    sz=OxmlElement('w:sz'); sz.set(qn('w:val'),'22'); rPr.append(sz)
    b=OxmlElement('w:b'); rPr.append(b)
    r.append(rPr); t=OxmlElement('w:t'); t.text=text; r.append(t)
    h.append(r); paragraph._p.append(h)

def build(data, out_path):
    doc=Document(); base=doc.styles["Normal"]; base.font.name="Arial"; base.font.size=Pt(11)
    sec=doc.sections[0]
    sec.left_margin=Inches(1); sec.right_margin=Inches(1); sec.top_margin=Inches(0.8); sec.bottom_margin=Inches(0.8)

    def spacer(pts=6):
        p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(pts); p.paragraph_format.space_before=Pt(0); return p

    # TOP: logo + legend
    top=doc.add_table(rows=1,cols=2); lcell,rcell=top.rows[0].cells; set_w(lcell,4.3); set_w(rcell,2.2); fix_table(top,[4.3,2.2])
    lcell.paragraphs[0].add_run().add_picture(LOGO,width=Inches(2.0))
    borders(rcell); shade(rcell,"FFFFFF")
    t=rcell.paragraphs[0]; tr=t.add_run("LEGEND"); tr.bold=True; tr.font.size=Pt(8); tr.font.color.rgb=LABELGRAY
    l1=rcell.add_paragraph(); a=l1.add_run("Black"); a.bold=True; a.font.size=Pt(8.5); b=l1.add_run(" = what the client says"); b.font.size=Pt(8.5)
    l2=rcell.add_paragraph(); c=l2.add_run("Red"); c.bold=True; c.font.size=Pt(8.5); c.font.color.rgb=RED; d=l2.add_run(" = what the interviewer says"); d.font.size=Pt(8.5); d.font.color.rgb=RED
    for p in rcell.paragraphs: p.paragraph_format.space_after=Pt(1); p.paragraph_format.space_before=Pt(1)
    spacer(2)

    # Header — big shoot number, then client + IG (no date)
    shoot=str(data.get("shoot","")).strip()
    if shoot:
        sp=doc.add_paragraph(); sp.paragraph_format.space_after=Pt(3)
        sr=sp.add_run(shoot); sr.bold=True; sr.font.size=Pt(20)
    for lab,val in (("Client:",data.get("client","")),("IG:",data.get("ig",""))):
        if not str(val).strip(): continue
        p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(0)
        r=p.add_run((lab+" "+val).strip()); r.bold=True; r.font.size=Pt(13); r.font.highlight_color=WD_COLOR_INDEX.TURQUOISE
    spacer(6)

    # Instructions box
    itbl=doc.add_table(rows=2,cols=1); ih=itbl.rows[0].cells[0]; ib=itbl.rows[1].cells[0]; set_w(ih,6.5); set_w(ib,6.5); fix_table(itbl,[6.5])
    shade(ih,LABEL_FILL); borders(ih)
    hp=ih.paragraphs[0]; hp.paragraph_format.space_after=Pt(2); hp.paragraph_format.space_before=Pt(2)
    hr=hp.add_run("HOW TO REVIEW THIS DOC"); hr.bold=True; hr.font.size=Pt(10); hr.font.color.rgb=LABELGRAY
    borders(ib); shade(ib,NOTE_FILL)
    ib.paragraphs[0]._element.getparent().remove(ib.paragraphs[0]._element)
    def iline(runs,sa=3,bullet=False):
        p=ib.add_paragraph(); p.paragraph_format.space_after=Pt(sa); p.paragraph_format.space_before=Pt(2)
        if bullet: p.paragraph_format.left_indent=Inches(0.18)
        for tup in runs:
            txt,bold,color=tup[0],tup[1],tup[2]; hl=tup[3] if len(tup)>3 else None
            r=p.add_run(txt); r.bold=bold; r.font.size=Pt(10)
            if color is not None: r.font.color.rgb=color
            if hl is not None: r.font.highlight_color=hl
        return p
    iline([("Read each script. If you want to change any wording, ",False,None),("just edit it right here in the doc",True,None),(" — you don't have to ask us first. Then ",False,None),("highlight the video title",True,None),(" to tell us your call:",False,None)])
    iline([("● ",False,None),("Approved",True,BLACK,WD_COLOR_INDEX.BRIGHT_GREEN),("  — good to shoot. Any edits you made are final.",False,None)],bullet=True)
    iline([("● ",False,None),("Change it",True,BLACK,WD_COLOR_INDEX.YELLOW),("  — you like the idea but want something different (new topic, different angle, or a new spin). Tell us what in a comment and we'll rewrite it.",False,None)],bullet=True)
    iline([("● ",False,None),("Not this one",True,BLACK,WD_COLOR_INDEX.RED),("  — a no. Tell us why in a comment (and if you have a different idea, drop it there too), so we learn what to avoid.",False,None)],bullet=True)
    iline([("To leave a note, highlight the text and click the ",False,None),("comment button",True,None),(" (shown below).",False,None)],sa=4)
    iline([("Please edit this Word file right here in Drive",True,RED),(" — do ",False,None),("not",True,None),(" convert it to a Google Doc or make a copy. Keeping it as one file is what lets us send your revised version back to this same link.",False,None)],sa=4)
    imgp=ib.add_paragraph(); imgp.paragraph_format.space_after=Pt(4); imgp.add_run().add_picture(COMMENT,width=Inches(2.9))

    # field/box helpers
    def fields_grid(topic, fmt, fmt_link, text_hook):
        tbl=doc.add_table(rows=3,cols=2)
        data_rows=[("TOPIC",topic,None),("FORMAT",fmt,fmt_link),("TEXT HOOK",text_hook,None)]
        for i,(lab,val,link) in enumerate(data_rows):
            lc,vc=tbl.rows[i].cells; set_w(lc,1.25); set_w(vc,5.25); lc.vertical_alignment=WD_ALIGN_VERTICAL.CENTER
            shade(lc,LABEL_FILL); shade(vc,"FFFFFF"); borders(lc); borders(vc)
            lp=lc.paragraphs[0]; lp.paragraph_format.space_after=Pt(2); lp.paragraph_format.space_before=Pt(2)
            lr=lp.add_run(lab); lr.bold=True; lr.font.size=Pt(9.5); lr.font.color.rgb=LABELGRAY
            vp=vc.paragraphs[0]; vp.paragraph_format.space_after=Pt(2); vp.paragraph_format.space_before=Pt(2)
            if val: vr=vp.add_run(val); vr.font.size=Pt(11)
            if link:
                if val: vp.add_run("  ")
                add_hyperlink(vp, link, "HERE")
        fix_table(tbl,[1.25,5.25])

    def boxed_section(title, text=None, lines=None, blank=2):
        tbl=doc.add_table(rows=2,cols=1); hc=tbl.rows[0].cells[0]; bc=tbl.rows[1].cells[0]; set_w(hc,6.5); set_w(bc,6.5)
        shade(hc,LABEL_FILL); borders(hc)
        hp=hc.paragraphs[0]; hp.paragraph_format.space_after=Pt(2); hp.paragraph_format.space_before=Pt(2)
        hr=hp.add_run(title); hr.bold=True; hr.font.size=Pt(9.5); hr.font.color.rgb=LABELGRAY
        borders(bc); bc.paragraphs[0].paragraph_format.space_after=Pt(3)
        if lines:
            first=True
            for item in lines:
                p=bc.paragraphs[0] if first else bc.add_paragraph(); first=False
                marker=item.get("marker","●")  # "" = plain line (no bullet); e.g. "1." for numbered
                p.paragraph_format.left_indent=Inches(0.18 if marker else 0); p.paragraph_format.space_after=Pt(4)
                if marker:
                    bl=p.add_run(marker+"  "); bl.font.size=Pt(11); bl.font.color.rgb=BLACK
                if "runs" in item:
                    for rn in item["runs"]:
                        col=RED if rn.get("who")=="interviewer" else BLACK
                        r=p.add_run(rn.get("t","")); r.font.size=Pt(11); r.font.color.rgb=col
                else:
                    col=RED if item.get("who")=="interviewer" else BLACK
                    r=p.add_run(item.get("t","")); r.font.size=Pt(11); r.font.color.rgb=col
        elif text:
            bc.paragraphs[0].add_run(text).font.size=Pt(11)
        else:
            for _ in range(blank-1): bc.add_paragraph()
        fix_table(tbl,[6.5])

    # one video per page
    for i, v in enumerate(data.get("videos",[]), 1):
        vp=doc.add_paragraph(); vp.paragraph_format.page_break_before=True
        vp.paragraph_format.space_before=Pt(2); vp.paragraph_format.space_after=Pt(6)
        vr=vp.add_run(f"Video {i}"); vr.bold=True; vr.font.size=Pt(16)
        if str(v.get("verdict","")).lower()=="approved":
            vr.font.highlight_color=WD_COLOR_INDEX.BRIGHT_GREEN  # keep client-approved videos green
        fields_grid(v.get("topic",""), v.get("format",""), v.get("format_link"), v.get("text_hook",""))
        spacer(4)
        boxed_section("EDITOR NOTES  —  for the editor: what to shoot / how to cut (not spoken)", text=(v.get("editor_notes") or None), blank=2)
        spacer(4)
        boxed_section("SCRIPT", lines=(v.get("script") or None), blank=6)

    doc.save(out_path)

if __name__ == "__main__":
    data_path, out_path = sys.argv[1], sys.argv[2]
    with open(data_path, encoding="utf-8") as f:
        data = json.load(f)
    build(data, out_path)
    print("saved", out_path)
