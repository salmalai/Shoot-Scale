#!/usr/bin/env python3
"""Build a client's Idea Bank as a Word .docx (kept under 500 characters of text).

The Idea Bank is a LIVING Word document (never a .md): the PM pastes client-sent
"make a video like this" ideas under "Open ideas," and /produce moves each idea to
"Used ideas" once it's scripted so nothing is ever pulled twice.

Usage:
    python3 build_idea_bank.py "<Client Name>" "<output path>/Idea-Bank.docx"

Needs python-docx:  pip install python-docx --break-system-packages
Create-if-missing:  the caller must NOT run this if an Idea-Bank.docx already exists.
"""
import sys
from docx import Document
from docx.shared import Pt, RGBColor

CYAN = RGBColor(0x00, 0x8C, 0xBA)
GREY = RGBColor(0x66, 0x66, 0x66)


def paragraphs(client):
    """Return the doc's lines as (text, kind) — kept under 500 chars total."""
    return [
        (f"{client} — Idea Bank", "title"),
        ("Paste new ideas under Open ideas (link + a note). /produce uses each once, "
         "then moves it to Used ideas — never reused.", "body"),
        ("Open ideas (not yet used)", "head"),
        ("(none yet — paste them here)", "muted"),
        ("Used ideas (already scripted — do not reuse)", "head"),
        ("(none yet)", "muted"),
        ("Keep as .docx — don’t convert to a Google Doc.", "foot"),
    ]


def build(client: str, out_path: str) -> int:
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    total = 0
    for text, kind in paragraphs(client):
        total += len(text)
        p = doc.add_paragraph()
        r = p.add_run(text)
        if kind == "title":
            r.bold = True; r.font.size = Pt(20); r.font.color.rgb = CYAN
        elif kind == "head":
            r.bold = True; r.font.size = Pt(14)
        elif kind == "muted":
            r.italic = True; r.font.color.rgb = GREY
        elif kind == "foot":
            r.italic = True; r.font.size = Pt(9); r.font.color.rgb = GREY
    doc.save(out_path)
    return total


if __name__ == "__main__":
    if len(sys.argv) != 3:
        sys.exit("Usage: build_idea_bank.py \"<Client Name>\" \"<output path>/Idea-Bank.docx\"")
    n = build(sys.argv[1], sys.argv[2])
    print(f"wrote {sys.argv[2]} ({n} chars of text)")
